"""
This is my final project. I chose the code that generates a recipe book from taco recipes including a random taco picture
"""
import requests # import request library
import sys # import sys library
from PIL import Image, ImageFont, ImageDraw # import libraries for font and image
from docx import Document # to create word document
from docx.shared import Inches
from pyunsplash import PyUnsplash # import from Unsplash website

# api for unsplash
API_KEY = "6fe959d9f7971906d0cb9fa31671d876772c1b841405f91b914f2b55bd838257"
# taco api url
URL = "https://taco-1150.herokuapp.com/random/?full_taco=true"
FONT_FILE = "DejaVuSans.ttf"


# this function is to search for image
def search_image(tag):
    '''
    search and download one image with this tag
    '''

    # authenticate unsplash api
    try:
        py_un = PyUnsplash(api_key=API_KEY)
        search = py_un.search(type_='photos', query=tag)
        photos = [i.link_download for i in search.entries]
        image_request = requests.get(photos[0], verify=False)
        with open("taco.jpg", 'wb') as img:
            img.write(image_request.content)
        return photos[0]
    except:
        print("Error in search_image")
        sys.exit(1)


def resize_image():
    '''
    this function resize the image and write Random Taco CookBook
    '''
    try:
        im = Image.open('taco.jpg')
    except IOError:
        print("Error opening image")
        sys.exit(1)
    width, height = im.size
    ratio = height / width
    # set image to max 800 in width and height avoiding all logical errors
    if (width > height):
        width = 800 # set width size to 800
        height = int(ratio * 800) # multiply height ratio by 800
    else:
        height = 800 # set height to 800
        width = int(800 / ratio) # multiply width ratio by 800
    im = im.resize((width, height), Image.ANTIALIAS)
    # adding font to taco image
    font = ImageFont.truetype(FONT_FILE, 40) # setting font size
    draw = ImageDraw.Draw(im) # drawing function
    draw.text((0, 0), "Random Taco Cookbook", (255, 255, 255), font=font) # choosing pixels to place the font
    im.save('taco_small.jpg', format='JPEG') # saving new modified image to taco small


def make_request():
    '''
    this function returns data from taco api
    '''
    try:
        data = requests.get(URL, verify=False) # returning data from taco API while verifying
        data = data.json()
    except:
        print("Error requesting taco API") # statement prints out in case of error while requesting from API
        sys.exit(1)
    return data


# used a function to create & customize document
def create_document(data, image_name, photo_info):
    document = Document()
    # naming and adding heading
    document.add_heading("Random Taco Cookbook", 0)
    # adding a picture
    document.add_picture(image_name, width=Inches(5.5))
    # adding and naming heading for credits
    document.add_heading('Credits')
    # adding lines under credits
    document.add_paragraph('Taco image: ' + photo_info, style='List Bullet') # adds unsplash link, styles it with bullet points
    document.add_paragraph('Tacos from: ' + URL, style='List Bullet') # adds taco recipes API, styles it with bullet points
    document.add_paragraph('Code by: Khalil', style='List Bullet') # adds my name, styles it with bullet points

    # this loops through the recipes in URL (taco recipes) and creates a document by listing them
    for recipe in data: # loops each recipe in taco recipes
        document.add_page_break() # adds a page break
        names = ', '.join(recipe[name]['name'] for name in recipe) # creates a headline for each recipe
        document.add_heading(names, 0)

        for rec in recipe:  # this loops through recipes list created from loop above, creates a list
            document.add_heading(recipe[rec]['name'], level=1) # creates a heading for recipe list
            document.add_paragraph(recipe[rec]['recipe']) # adds list of recipe
    # after looping, save the data in a word document
    document.save("recipes.docx")


"""
p = search_image('taco') will search in unsplash for one image that has 'taco' as tag. 
resize_image() will resize that image (p), and that loop will store 5 recipes taken from that Taco API.
create_document will use the resized image and the five recipes to create the book of recipes
"""

def main():
    p = search_image('taco') # search for image with taco in tag
    resize_image() # use this function to resize the image
    recipes = [] # create an empty list
    for i in range(5): # create a sequence of recipes and print them in that sequence = 5
        recipes.append(make_request()) # appending recipes
    create_document(recipes, 'taco_small.jpg', p) # create the book of recipes

if __name__ == '__main__':
    main()

"""
The End...
"""
